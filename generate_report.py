#!/usr/bin/env python3
"""
generate_report.py  –  CNG 436 Assignment 2 report (full code edition).
Run from project root:  python3 generate_report.py
"""

import os
from PIL import Image
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE = os.path.dirname(os.path.abspath(__file__))
OUT  = os.path.join(BASE, "CNG436_Assignment2_Report (1).docx")

# ── helpers ───────────────────────────────────────────────────────────────────

def convert_webp(src, dst):
    if not os.path.exists(src):
        return False
    if not os.path.exists(dst):
        Image.open(src).save(dst, "PNG")
    return True

def set_cell_bg(cell, hex_color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    tcPr.append(shd)

def heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return p

def body(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(11)
    return p

def placeholder(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(f"[ {text} ]")
    run.italic = True
    run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)
    run.font.size = Pt(11)
    return p

def add_image(doc, path, width_cm=14, caption=None):
    if os.path.exists(path):
        doc.add_picture(path, width=Cm(width_cm))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        ph = doc.add_paragraph()
        ph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = ph.add_run(f"[ INSERT IMAGE: {os.path.basename(path)} ]")
        run.bold = True
        run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)
        run.font.size = Pt(12)
    if caption:
        cp = doc.add_paragraph(caption)
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cp.runs[0].italic = True
        cp.runs[0].font.size = Pt(10)

def add_code(doc, code_text, caption=None):
    if caption:
        cp = doc.add_paragraph(caption)
        cp.runs[0].bold = True
        cp.runs[0].font.size = Pt(10)
    p = doc.add_paragraph()
    run = p.add_run(code_text)
    run.font.name = 'Courier New'
    run.font.size = Pt(8.5)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  'F2F2F2')
    pPr.append(shd)

def page_break(doc):
    doc.add_page_break()

# ── validation table data ─────────────────────────────────────────────────────

CONFIGS = [
    ('MM1L_light',    1.0, 3.0, 10, 0.3333,
     [('E[W]',       0.166610, 0.165000, 0.97),
      ('E[W|W>0]',   0.499836, 0.497556, 0.46),
      ('U',          0.333330, 0.333402, 0.02),
      ('E[Nq]',      0.166608, 0.164849, 1.06),
      ('lambda_eff', 0.999989, 0.999071, 0.09),
      ('E[T]',       0.499944, 0.498709, 0.25)]),
    ('MM1L_medium',   2.0, 3.0, 10, 0.6667,
     [('E[W]',       0.607841, 0.615584, 1.27),
      ('E[W|W>0]',   0.917126, 0.927884, 1.17),
      ('U',          0.662768, 0.664388, 0.24),
      ('E[Nq]',      1.208573, 1.226224, 1.46),
      ('lambda_eff', 1.988304, 1.991939, 0.18),
      ('E[T]',       0.941175, 0.949119, 0.84)]),
    ('MM1L_heavy',    2.5, 3.0, 10, 0.8333,
     [('E[W]',       1.024621, 1.037003, 1.21),
      ('E[W|W>0]',   1.269016, 1.290070, 1.66),
      ('U',          0.807414, 0.809175, 0.22),
      ('E[Nq]',      2.481878, 2.513746, 1.28),
      ('lambda_eff', 2.422241, 2.424061, 0.08),
      ('E[T]',       1.357954, 1.370810, 0.95)]),
    ('MM1L_overload', 4.0, 3.0, 10, 1.3333,
     [('E[W]',       2.198913, 2.189871, 0.41),
      ('E[W|W>0]',   2.231718, 2.234679, 0.13),
      ('U',          0.985301, 0.985326, 0.00),
      ('E[Nq]',      6.499773, 6.482018, 0.27),
      ('lambda_eff', 2.955902, 2.960081, 0.14),
      ('E[T]',       2.532247, 2.522746, 0.38)]),
]

def add_validation_table(doc, cfg_name, lam, mu, L, rho, rows):
    p = doc.add_paragraph()
    run = p.add_run(f"Config: {cfg_name}   |   lambda={lam}, mu={mu}, L={L}, rho={rho:.4f}")
    run.bold = True
    run.font.size = Pt(10)

    tbl = doc.add_table(rows=1 + len(rows), cols=4)
    tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(['Metric', 'Analytical', 'Simulation', '|Diff| %']):
        cell = tbl.rows[0].cells[i]
        cell.text = h
        cell.paragraphs[0].runs[0].bold = True
        set_cell_bg(cell, '1F497D')
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    for r_idx, (metric, ana, sim, diff) in enumerate(rows):
        cells = tbl.rows[r_idx + 1].cells
        cells[0].text = metric
        cells[1].text = f"{ana:.6f}"
        cells[2].text = f"{sim:.6f}"
        cells[3].text = f"{diff:.2f}%"
        bg = 'DDEEFF' if r_idx % 2 == 0 else 'FFFFFF'
        for c in cells:
            set_cell_bg(c, bg)
    doc.add_paragraph()

# ── full source code strings ──────────────────────────────────────────────────

JOB_MSG = """\
// Job.msg  –  message definition for the M/M/1/L simulator
message Job
{
    double arrivalTime;   // simTime() when Source created this job
}"""

SOURCE_H = """\
#ifndef SOURCE_H
#define SOURCE_H

#include <omnetpp.h>
using namespace omnetpp;

// Source: generates Poisson arrivals (exponential inter-arrival times).
// Inter-arrival time drawn via inverse transform: -ln(U) / lambda.
class Source : public cSimpleModule
{
private:
    double   lambda;
    cMessage *arrivalEvent;
    long     totalGenerated;

protected:
    virtual void initialize()            override;
    virtual void handleMessage(cMessage*) override;
    virtual void finish()                override;
};

#endif"""

SOURCE_CC = """\
#include "Source.h"
#include "../msg/Job_m.h"

Define_Module(Source);

void Source::initialize()
{
    lambda         = par("lambda").doubleValue();
    totalGenerated = 0;

    arrivalEvent = new cMessage("arrival");
    double ia = -std::log(dblrand()) / lambda;   // Exp(lambda)
    scheduleAt(simTime() + ia, arrivalEvent);

    EV << "[Source] initialized. lambda=" << lambda << "/s" << endl;
}

void Source::handleMessage(cMessage *msg)
{
    ASSERT(msg == arrivalEvent);

    Job *job = new Job("job");
    job->setArrivalTime(simTime().dbl());   // timestamp for E[T] and E[W]
    send(job, "out");
    totalGenerated++;

    double ia = -std::log(dblrand()) / lambda;
    scheduleAt(simTime() + ia, arrivalEvent);
}

void Source::finish()
{
    recordScalar("totalGenerated", totalGenerated);
    cancelAndDelete(arrivalEvent);
    arrivalEvent = nullptr;
    EV << "[Source] finish(): totalGenerated=" << totalGenerated << endl;
}"""

SOURCE_NED = """\
// Source.ned  –  Poisson arrival process module
simple Source
{
    parameters:
        double lambda = default(1.0);   // arrival rate (jobs/s)
        @display("i=block/source");
    gates:
        output out;
}"""

SERVER_H = """\
#ifndef SERVER_H
#define SERVER_H

#include <omnetpp.h>
#include <queue>
#include "../msg/Job_m.h"
using namespace omnetpp;

// Server: M/M/1/L single-server FIFO queue with finite system capacity L.
// Arrivals accepted if systemSize() < L, else blocked and dropped.
// Service times: Exp(mu) via inverse transform: -ln(U)/mu.
class Server : public cSimpleModule
{
private:
    int    capacity;        // L: max customers in system (queue + server)
    double mu;              // service rate (jobs/s)

    std::queue<Job *> jobQueue;
    bool              serverBusy;
    Job              *jobInService;
    cMessage         *serviceEndEvent;

    // Post-warmup counters (reset by warmupResetEvent at end of warmup period)
    long   totalArrived;
    long   totalDropped;
    long   totalServed;
    long   countWaited;     // jobs that had W > 0
    double sumWaitGT0;      // sum of positive wait times

    cMessage    *warmupResetEvent;
    simsignal_t  waitingTimeSignal;
    simsignal_t  queueLengthSignal;
    simsignal_t  utilisationSignal;
    simsignal_t  systemLengthSignal;

protected:
    virtual void initialize()             override;
    virtual void handleMessage(cMessage*) override;
    virtual void finish()                 override;

private:
    void startService(Job *job);
    void endService();
    void resetCounters();
    int  systemSize() const { return (int)jobQueue.size() + (serverBusy ? 1 : 0); }
};

#endif"""

SERVER_CC = """\
#include "Server.h"

Define_Module(Server);

void Server::initialize()
{
    capacity        = par("capacity");
    mu              = par("mu").doubleValue();
    serverBusy      = false;
    jobInService    = nullptr;
    serviceEndEvent = new cMessage("serviceEnd");

    resetCounters();

    waitingTimeSignal  = registerSignal("waitingTime");
    queueLengthSignal  = registerSignal("queueLength");
    utilisationSignal  = registerSignal("utilisation");
    systemLengthSignal = registerSignal("systemLength");

    emit(queueLengthSignal,  0L);
    emit(utilisationSignal,  0.0);
    emit(systemLengthSignal, 0L);

    // Schedule counter reset at warmup end so scalars cover only steady-state.
    SimTime wp = getSimulation()->getWarmupPeriod();
    if (wp > SimTime::ZERO) {
        warmupResetEvent = new cMessage("warmupReset");
        scheduleAt(wp, warmupResetEvent);
    } else {
        warmupResetEvent = nullptr;
    }
    EV << "[Server] initialized. capacity=" << capacity
       << " mu=" << mu << "/s" << endl;
}

void Server::resetCounters()
{
    totalArrived = 0;
    totalDropped = 0;
    totalServed  = 0;
    countWaited  = 0;
    sumWaitGT0   = 0.0;
}

void Server::handleMessage(cMessage *msg)
{
    // 1) Warmup reset: discard transient statistics
    if (msg == warmupResetEvent) {
        resetCounters();
        EV << "[Server] Warmup ended. Counters reset." << endl;
        delete warmupResetEvent;
        warmupResetEvent = nullptr;
        return;
    }

    // 2) Service completion
    if (msg == serviceEndEvent) {
        endService();
        return;
    }

    // 3) New arrival
    Job *job = check_and_cast<Job *>(msg);
    totalArrived++;

    if (systemSize() >= capacity) {    // system full -> BLOCK (implements P_L)
        totalDropped++;
        EV << "[Server] DROP (system full, L=" << capacity
           << "). drops=" << totalDropped << endl;
        delete job;
        return;
    }

    if (!serverBusy) {
        startService(job);             // server idle -> go straight to service
    } else {
        jobQueue.push(job);            // server busy -> join queue
        emit(queueLengthSignal,  (long)jobQueue.size());
        emit(systemLengthSignal, (long)systemSize());
        EV << "[Server] Queued. queue=" << jobQueue.size() << endl;
    }
}

void Server::startService(Job *job)
{
    serverBusy   = true;
    jobInService = job;

    double waitTime = simTime().dbl() - job->getArrivalTime();
    emit(waitingTimeSignal,  waitTime);          // feeds E[W] statistic
    emit(utilisationSignal,  1.0);
    emit(systemLengthSignal, (long)systemSize());

    if (waitTime > 0.0) {               // accumulate E[W|W>0] numerator
        countWaited++;
        sumWaitGT0 += waitTime;
    }

    double svcTime = -std::log(dblrand()) / mu; // Exp(mu) service time
    scheduleAt(simTime() + svcTime, serviceEndEvent);

    EV << "[Server] Service started. wait=" << waitTime
       << "s svc=" << svcTime << "s" << endl;
}

void Server::endService()
{
    ASSERT(jobInService != nullptr);
    Job *job = jobInService;
    jobInService = nullptr;

    send(job, "out");           // forward to Sink for response-time recording
    totalServed++;
    EV << "[Server] Service done. total served=" << totalServed << endl;

    if (!jobQueue.empty()) {    // pick next job from queue (FIFO)
        Job *next = jobQueue.front();
        jobQueue.pop();
        emit(queueLengthSignal, (long)jobQueue.size());
        startService(next);
    } else {                    // queue empty -> server goes idle
        serverBusy = false;
        emit(utilisationSignal,  0.0);
        emit(queueLengthSignal,  0L);
        emit(systemLengthSignal, 0L);
    }
}

void Server::finish()
{
    SimTime wp          = getSimulation()->getWarmupPeriod();
    double  simDuration = (simTime() - wp).dbl();

    recordScalar("totalArrived", totalArrived);
    recordScalar("totalDropped", totalDropped);
    recordScalar("totalServed",  totalServed);

    double lambdaEff    = (simDuration > 0) ? (double)totalServed / simDuration : 0.0;
    double blockingProb = (totalArrived > 0) ? (double)totalDropped / totalArrived : 0.0;
    double waitGT0Mean  = (countWaited > 0)  ? sumWaitGT0 / countWaited : 0.0;

    recordScalar("lambda:eff",          lambdaEff);     // approx lambda*(1-P_L)
    recordScalar("blocking:prob",       blockingProb);  // approx P_L
    recordScalar("waitingTimeGT0:mean", waitGT0Mean);   // approx E[W|W>0]

    EV << "[Server] finish(): arrived=" << totalArrived
       << " served=" << totalServed << " dropped=" << totalDropped
       << " lambdaEff=" << lambdaEff
       << " E[W|W>0]=" << waitGT0Mean << endl;

    cancelAndDelete(serviceEndEvent);
    serviceEndEvent = nullptr;
    if (warmupResetEvent) { cancelAndDelete(warmupResetEvent); warmupResetEvent = nullptr; }
    if (jobInService)     { delete jobInService; jobInService = nullptr; }
    while (!jobQueue.empty()) { delete jobQueue.front(); jobQueue.pop(); }
}"""

SERVER_NED = """\
// Server.ned  –  M/M/1/L queue module definition
simple Server
{
    parameters:
        int    capacity = default(10);  // L: max customers in system
        double mu       = default(3.0); // service rate (jobs/s)

        @display("i=block/queue");

        @signal[waitingTime](type=double);
        @signal[queueLength](type=long);
        @signal[utilisation](type=double);
        @signal[systemLength](type=long);

        @statistic[waitingTime](title="waiting time";
            source=waitingTime; unit=s; record=mean,vector; checkSignals=false);
        @statistic[queueLength](title="queue length";
            source=queueLength; record=mean,timeavg,vector; checkSignals=false);
        @statistic[utilisation](title="utilisation";
            source=utilisation; record=mean,timeavg; checkSignals=false);
        @statistic[systemLength](title="system length";
            source=systemLength; record=mean,timeavg,vector; checkSignals=false);
    gates:
        input  in;
        output out;
}"""

SINK_H = """\
#ifndef SINK_H
#define SINK_H

#include <omnetpp.h>
#include "../msg/Job_m.h"
using namespace omnetpp;

// Sink: absorbs completed jobs, records end-to-end response time E[T].
class Sink : public cSimpleModule
{
private:
    long        totalReceived;
    simsignal_t responseTimeSignal;

protected:
    virtual void initialize()             override;
    virtual void handleMessage(cMessage*) override;
    virtual void finish()                 override;
};

#endif"""

SINK_CC = """\
#include "Sink.h"

Define_Module(Sink);

void Sink::initialize()
{
    totalReceived      = 0;
    responseTimeSignal = registerSignal("responseTime");
    EV << "[Sink] initialized." << endl;
}

void Sink::handleMessage(cMessage *msg)
{
    Job *job = check_and_cast<Job *>(msg);

    // Response time = sojourn time = waiting + service  (Little's Law: E[N] = lambda_eff * E[T])
    double responseTime = simTime().dbl() - job->getArrivalTime();
    emit(responseTimeSignal, responseTime);
    totalReceived++;

    EV << "[Sink] Job received. responseTime=" << responseTime
       << "s  total=" << totalReceived << endl;
    delete job;
}

void Sink::finish()
{
    recordScalar("totalReceived", totalReceived);
    EV << "[Sink] finish(): totalReceived=" << totalReceived << endl;
}"""

SINK_NED = """\
// Sink.ned  –  job absorber module definition
simple Sink
{
    parameters:
        @display("i=block/sink");
        @signal[responseTime](type=double);
        @statistic[responseTime](title="response time";
            source=responseTime; unit=s; record=mean,vector; checkSignals=false);
    gates:
        input in;
}"""

MM1L_NED = """\
// MM1L.ned  –  M/M/1/L finite-capacity queueing system (network definition)
// Topology: Source --> Server (M/M/1/L queue) --> Sink
network MM1L
{
    parameters:
        @display("bgb=600,300");
    submodules:
        source : Source { @display("p=100,150"); }
        server : Server { @display("p=300,150"); }
        sink   : Sink   { @display("p=500,150"); }
    connections:
        source.out --> server.in;
        server.out --> sink.in;
}"""

OMNETPP_INI = """\
[General]
network          = MM1L
ned-path         = ../src
sim-time-limit   = 100000s
warmup-period    = 1000s
repeat           = 1
seed-set         = ${repetition}
result-dir       = results
num-rngs         = 2
rng-class        = "cMersenneTwister"

**.scalar-recording = true
**.vector-recording = false

# Fixed system parameters
**.server.capacity = 10      # L: finite system capacity
**.server.mu       = 3.0     # mu: service rate (jobs/s)

[Config MM1L_light]
description = "M/M/1/L: lambda=1, mu=3, L=10  (light load, rho=0.333)"
**.source.lambda = 1.0

[Config MM1L_medium]
description = "M/M/1/L: lambda=2, mu=3, L=10  (medium load, rho=0.667)"
**.source.lambda = 2.0

[Config MM1L_heavy]
description = "M/M/1/L: lambda=2.5, mu=3, L=10  (heavy load, rho=0.833)"
**.source.lambda = 2.5

[Config MM1L_overload]
description = "M/M/1/L: lambda=4, mu=3, L=10  (overload, rho=1.333)"
**.source.lambda = 4.0"""

VALIDATE_PY = """\
def mm1l_analytical(lam, mu, L):
    \"\"\"Closed-form M/M/1/L performance metrics (from balance equations).\"\"\"
    rho = lam / mu

    if abs(rho - 1.0) < 1e-10:          # special case rho = 1
        P0  = 1.0 / (L + 1)
        P_L = P0
        E_N = L / 2.0
    else:
        denom = 1.0 - rho ** (L + 1)
        P0    = (1.0 - rho) / denom                            # P0 formula
        P_L   = P0 * rho ** L                                  # P_n = rho^n * P0
        E_N   = rho/(1-rho) - (L+1)*rho**(L+1)/denom          # E[N] formula

    lam_eff = lam * (1.0 - P_L)         # throughput: lambda_eff = lambda*(1-P_L)
    U       = 1.0 - P0                  # utilisation: U = 1 - P0
    E_Nq    = E_N - U                   # E[Nq] = E[N] - U
    E_T     = E_N   / lam_eff           # Little's Law: E[T] = E[N] / lambda_eff
    E_W     = E_Nq  / lam_eff           # Little's Law: E[W] = E[Nq] / lambda_eff
    E_W_gt0 = E_W   / (1.0 - P0)       # E[W|W>0] = E[W] / P(W>0) = E[W] / (1-P0)

    return {
        'rho': rho, 'P0': P0, 'P_L': P_L,
        'lambda_eff': lam_eff, 'U': U,
        'E[N]': E_N, 'E[Nq]': E_Nq,
        'E[T]': E_T, 'E[W]': E_W, 'E[W|W>0]': E_W_gt0,
    }"""

# ── main build ────────────────────────────────────────────────────────────────

def build():
    imgs = {}
    for i in [1, 2, 3]:
        src = os.path.join(BASE, f"{i}.webp")
        dst = os.path.join(BASE, f"hw_{i}.png")
        convert_webp(src, dst)
        imgs[i] = dst

    doc = Document()

    for section in doc.sections:
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin   = Cm(3.0)
        section.right_margin  = Cm(2.5)

    # ── Title page ────────────────────────────────────────────────────────────
    for _ in range(6):
        doc.add_paragraph()

    for txt, sz, bold in [
        ("CNG 436 – Wireless Communication and Networks", 16, True),
        ("Assignment 2: Modelling and Simulation",        14, True),
        ("",                                              12, False),
        ("Spring 2025–2026",                              12, False),
        ("Middle East Technical University – Northern Cyprus Campus", 12, False),
        ("",                                              12, False),
        ("Author: Alp Kaan Ozgul",                       12, False),
        ("Submission Date: 20 May 2026",                  12, False),
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(txt)
        r.font.size = Pt(sz)
        r.bold = bold

    page_break(doc)

    # ── 1. Introduction ───────────────────────────────────────────────────────
    heading(doc, "1. Introduction")
    body(doc,
         "This report presents the design, implementation, and validation of discrete-event "
         "simulation models for finite-capacity queueing systems developed in OMNeT++ as "
         "part of CNG 436 Assignment 2. Two models are studied: the M/M/1/L system "
         "(single server, finite capacity L) and the M/M/c/L system (c parallel servers, "
         "finite capacity L). For each model an analytical solution is derived from balance "
         "equations and compared against simulation results to confirm accuracy within 5%.")
    page_break(doc)

    # ── 2. Question 1 – M/M/1/L Simulation ───────────────────────────────────
    heading(doc, "2. Question 1 – M/M/1/L Simulation (30%)")

    heading(doc, "2.1 System Description", level=2)
    body(doc,
         "The M/M/1/L queue has Poisson arrivals at rate lambda, a single exponential "
         "server at rate mu, and finite system capacity L (queue + server combined). "
         "Any arrival that finds the system full (n = L) is blocked and dropped. "
         "The simulator is built from scratch in OMNeT++ using three modules: "
         "Source, Server, and Sink, wired together in the MM1L network. "
         "Fixed parameters: mu = 3 jobs/s, L = 10. "
         "Four arrival rates are tested: lambda in {1.0, 2.0, 2.5, 4.0} jobs/s "
         "(rho in {0.333, 0.667, 0.833, 1.333}). "
         "Each run lasts 100,000 s with a 1,000 s warm-up period.")

    heading(doc, "2.2 Message Definition – Job.msg", level=2)
    body(doc,
         "Each job carries its creation timestamp so that waiting time and response time "
         "can be computed anywhere in the pipeline.")
    add_code(doc, JOB_MSG, caption="Listing 1 – msg/Job.msg")

    heading(doc, "2.3 Source Module", level=2)
    body(doc,
         "Source generates Poisson arrivals using the inverse-transform method: "
         "inter-arrival time = -ln(U)/lambda ~ Exp(lambda). Each job is timestamped.")

    add_code(doc, SOURCE_H,  caption="Listing 2 – src/Source.h")
    add_code(doc, SOURCE_NED, caption="Listing 3 – src/Source.ned")
    add_code(doc, SOURCE_CC, caption="Listing 4 – src/Source.cc")

    heading(doc, "2.4 Server Module", level=2)
    body(doc,
         "Server is the core of the M/M/1/L system. It enforces the capacity constraint "
         "(drops arrivals when full), draws exponential service times, maintains a FIFO "
         "queue, resets counters after the warm-up period, and records all statistics.")

    add_code(doc, SERVER_H,   caption="Listing 5 – src/Server.h")
    add_code(doc, SERVER_NED, caption="Listing 6 – src/Server.ned")
    add_code(doc, SERVER_CC,  caption="Listing 7 – src/Server.cc")

    heading(doc, "2.5 Sink Module", level=2)
    body(doc,
         "Sink absorbs completed jobs and records end-to-end response time "
         "E[T] = finish time - arrival time. By Little's Law: E[N] = lambda_eff * E[T].")

    add_code(doc, SINK_H,   caption="Listing 8 – src/Sink.h")
    add_code(doc, SINK_NED, caption="Listing 9 – src/Sink.ned")
    add_code(doc, SINK_CC,  caption="Listing 10 – src/Sink.cc")

    heading(doc, "2.6 Network Topology – MM1L.ned", level=2)
    add_code(doc, MM1L_NED, caption="Listing 11 – src/MM1L.ned")

    heading(doc, "2.7 Simulation Configuration – omnetpp.ini", level=2)
    add_code(doc, OMNETPP_INI, caption="Listing 12 – simulations/omnetpp.ini")

    page_break(doc)

    # ── 3. Question 2 – Analytical Model & Validation ─────────────────────────
    heading(doc, "3. Question 2 – Analytical Model & Validation (15%)")

    heading(doc, "3.1 State Diagram and Balance Equations (Handwritten)", level=2)
    body(doc,
         "The M/M/1/L system is modelled as a birth-death Markov chain with states "
         "n in {0, 1, ..., L}. The state diagram, balance equations, and derivation "
         "of all performance measures are provided in the handwritten pages below.")

    add_image(doc, imgs[1], width_cm=14,
              caption="Figure 1 – State diagram and balance equation derivation")
    add_image(doc, imgs[2], width_cm=14,
              caption="Figure 2 – General pattern Pn = rho^n * P0, finding P0, performance measures")
    add_image(doc, imgs[3], width_cm=14,
              caption="Figure 3 – E[W] and E[W|W>0] derivation via Little's Law")

    heading(doc, "3.2 Closed-Form Performance Measures", level=2)
    body(doc, "Summary of all derived analytical formulas for the M/M/1/L queue:")

    formulas = [
        ("P0 (rho != 1)",    "P0 = (1 - rho) / (1 - rho^(L+1))"),
        ("P0 (rho = 1)",     "P0 = 1 / (L + 1)"),
        ("State probability","Pn = rho^n * P0"),
        ("Blocking prob.",   "P_B = P_L = rho^L * P0"),
        ("Throughput",       "lambda_eff = lambda * (1 - P_L)"),
        ("Utilisation",      "U = 1 - P0"),
        ("Mean in system",   "E[N] = rho/(1-rho) - (L+1)*rho^(L+1) / (1-rho^(L+1))"),
        ("Mean queue length","E[Nq] = E[N] - U"),
        ("Response time",    "E[T] = E[N] / lambda_eff        (Little's Law)"),
        ("Waiting time",     "E[W] = E[Nq] / lambda_eff       (Little's Law)"),
        ("Wait | wait > 0",  "E[W|W>0] = E[W] / (1 - P0)"),
    ]
    tbl = doc.add_table(rows=1 + len(formulas), cols=2)
    tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(['Measure', 'Formula']):
        c = tbl.rows[0].cells[i]
        c.text = h
        c.paragraphs[0].runs[0].bold = True
        set_cell_bg(c, '1F497D')
        c.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    for i, (name, formula) in enumerate(formulas):
        row = tbl.rows[i + 1].cells
        row[0].text = name
        row[1].text = formula
        bg = 'DDEEFF' if i % 2 == 0 else 'FFFFFF'
        set_cell_bg(row[0], bg); set_cell_bg(row[1], bg)
    doc.add_paragraph()

    heading(doc, "3.3 Analytical Validation Script – validate.py", level=2)
    body(doc,
         "The Python script implements the analytical formulas above and reads the "
         "simulation .sca output files to produce a comparison table automatically.")
    add_code(doc, VALIDATE_PY, caption="Listing 13 – analysis/validate.py (analytical core)")

    heading(doc, "3.4 Validation Results", level=2)
    body(doc,
         "Each simulation was run for 100,000 s (warm-up 1,000 s). "
         "All discrepancies are well below the required 5% threshold. "
         "The maximum observed error is 1.66% (E[W|W>0] under heavy load).")

    for cfg_name, lam, mu, L, rho, rows in CONFIGS:
        add_validation_table(doc, cfg_name, lam, mu, L, rho, rows)

    body(doc,
         "Validation PASSED – all differences within 5% across all four "
         "configurations and all six metrics.")

    page_break(doc)

    # ── 4. Question 3 – M/M/c/L Simulation (template) ────────────────────────
    heading(doc, "4. Question 3 – M/M/c/L Simulation (30%)")

    heading(doc, "4.1 System Description", level=2)
    placeholder(doc, "Describe the M/M/c/L model: c parallel servers, capacity L, "
                     "same lambda and mu per server. State chosen parameter values.")

    heading(doc, "4.2 OMNeT++ Implementation", level=2)
    placeholder(doc, "Explain design changes relative to M/M/1/L: how c servers "
                     "are managed, queue discipline, blocking condition (system >= L).")

    heading(doc, "4.3 Module Code – MMcL Server", level=2)
    placeholder(doc, "Insert full source code: MMcLServer.h, MMcLServer.cc, "
                     "MMcLServer.ned, MMcL.ned, and updated omnetpp.ini configs.")

    heading(doc, "4.4 Simulation Configuration", level=2)
    placeholder(doc, "Show omnetpp.ini configs for M/M/c/L: lambda values, c=2, c=4, L.")

    page_break(doc)

    # ── 5. Question 4 – M/M/c/L Analytical + Validation (template) ───────────
    heading(doc, "5. Question 4 – M/M/c/L Analytical Model & Validation (15%)")

    heading(doc, "5.1 State Diagram and Balance Equations (Handwritten)", level=2)
    placeholder(doc, "Insert handwritten state diagram for M/M/c/L. Show separate "
                     "balance equations for n < c (multiple servers active) and "
                     "n >= c (all servers busy, queue growing).")

    heading(doc, "5.2 Closed-Form Performance Measures", level=2)
    placeholder(doc, "Insert table of M/M/c/L analytical formulas: P0, P_L, "
                     "lambda_eff, U, E[N], E[Nq], E[T], E[W], E[W|W>0].")

    heading(doc, "5.3 Validation Results – c = 2", level=2)
    placeholder(doc, "Insert comparison table for c=2 across various lambda values. "
                     "All |Diff|% must be < 5%.")

    heading(doc, "5.4 Validation Results – c = 4", level=2)
    placeholder(doc, "Insert comparison table for c=4 across various lambda values. "
                     "All |Diff|% must be < 5%.")

    page_break(doc)

    # ── 6. Conclusion ─────────────────────────────────────────────────────────
    heading(doc, "6. Conclusion")
    body(doc,
         "The M/M/1/L simulation was successfully validated against the analytical "
         "closed-form solution. All six performance metrics (E[W], E[W|W>0], U, "
         "E[Nq], lambda_eff, E[T]) were within 1.66% of theoretical values across "
         "four traffic loads (rho in {0.333, 0.667, 0.833, 1.333}). A 100,000 s "
         "simulation with 1,000 s warm-up is sufficient for steady-state accuracy.")
    placeholder(doc, "Add conclusion for M/M/c/L once Q3/Q4 are complete.")

    page_break(doc)

    # ── 7. References ─────────────────────────────────────────────────────────
    heading(doc, "7. References")
    refs = [
        '[1] N. van Foreest, "Simulating Queueing Networks with OMNeT++," '
        'https://omnetpp.org/download-items/Queues.html, 2003.',
        '[2] OMNeT++ Community, OMNeT++ Discrete Event Simulation System User Manual, 2024. '
        'Available: https://doc.omnetpp.org/omnetpp/manual/',
        '[3] INET Framework Developers, "INET Framework Users Guide: Queueing Model," '
        'https://inet.omnetpp.org/docs/users-guide/ch-queueing.html, 2024.',
        '[4] INET Framework Developers, "INET Framework Documentation," 2024. '
        'Available: https://inet.omnetpp.org',
    ]
    for ref in refs:
        p = doc.add_paragraph(ref, style='List Number')
        p.runs[0].font.size = Pt(10)

    doc.save(OUT)
    print(f"Report saved -> {OUT}")

if __name__ == '__main__':
    build()
